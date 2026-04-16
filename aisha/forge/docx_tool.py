"""DOCX generation.

Minimal surface: build a .docx from a structured outline (title + sections,
each section a heading plus paragraphs and/or bullets). Upload paths reuse
``pptx_tool.upload_to_slack`` — no duplicated Slack plumbing here.
"""
from __future__ import annotations

import logging
import tempfile
import uuid
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)


def generate_docx(
    title: str,
    sections: list[dict],
    *,
    subtitle: str = "",
    out_path: Optional[Path] = None,
) -> Path:
    """Build a .docx from a structured outline.

    ``sections`` is a list of ``{"heading": str, "paragraphs": [str, ...],
    "bullets": [str, ...]}`` dicts — either ``paragraphs`` or ``bullets`` (or
    both) per section. Returns the path to the generated file.
    """
    from docx import Document  # lazy import — keeps chat boot fast

    doc = Document()
    doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True

    for item in sections:
        heading = (item.get("heading") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        for para in item.get("paragraphs") or []:
            text = str(para).strip()
            if text:
                doc.add_paragraph(text)
        for bullet in item.get("bullets") or []:
            text = str(bullet).strip()
            if text:
                doc.add_paragraph(text, style="List Bullet")

    if out_path is None:
        name = f"aisha_{uuid.uuid4().hex[:8]}.docx"
        out_path = Path(tempfile.gettempdir()) / name
    doc.save(str(out_path))
    log.info("docx: generated %s (%d sections)", out_path, len(sections))
    return out_path
